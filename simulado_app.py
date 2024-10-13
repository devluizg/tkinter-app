import tkinter as tk
from tkinter import messagebox, filedialog
from tkinter.scrolledtext import ScrolledText
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from PIL import Image, ImageTk, ImageGrab
import base64
from io import BytesIO
import uuid
import re
import leitor_cartao

class SimuladoApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Gerador de Simulados")
        self.root.configure(bg="#333333")  # Cor de fundo externa (cinza mais escuro)

        self.questoes = []
        self.current_frame = None
        self.editing_index = None  # Para rastrear se estamos editando uma questão

        self.show_entry_frame()  # Mostra o quadro de entrada primeiro

    def show_entry_frame(self):
        """Exibe o quadro de entrada para adicionar/editar questões."""
        if self.current_frame:
            self.current_frame.destroy()

        self.current_frame = tk.Frame(self.root, bg="#333333")
        self.current_frame.pack(pady=20, expand=True, fill=tk.BOTH)

        # Configurar colunas e linhas para redimensionar com a janela
        self.current_frame.grid_columnconfigure(0, weight=1)
        self.current_frame.grid_columnconfigure(1, weight=1)
        self.current_frame.grid_rowconfigure(0, weight=1)
        self.current_frame.grid_rowconfigure(1, weight=0)
        self.current_frame.grid_rowconfigure(2, weight=0)
        self.current_frame.grid_rowconfigure(3, weight=0)
        self.current_frame.grid_rowconfigure(4, weight=0)
        self.current_frame.grid_rowconfigure(5, weight=0)

        self.label_enunciado = tk.Label(self.current_frame, text="Enunciado:", bg="#333333", fg="#FFFFFF")
        self.label_enunciado.grid(row=0, column=0, padx=5, pady=5, sticky="nw")

        self.text_enunciado = ScrolledText(self.current_frame, width=60, height=10, bg="#555555", fg="#FFFFFF", insertbackground="#FFFFFF")
        self.text_enunciado.grid(row=0, column=1, padx=5, pady=5, sticky="nsew")
        self.text_enunciado.bind("<Control-v>", self.paste_text_with_image)

        self.label_disciplina = tk.Label(self.current_frame, text="Disciplina:", bg="#333333", fg="#FFFFFF")
        self.label_disciplina.grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.entry_disciplina = tk.Entry(self.current_frame, width=30, bg="#555555", fg="#FFFFFF", insertbackground="#FFFFFF")
        self.entry_disciplina.grid(row=1, column=1, padx=5, pady=5, sticky="ew")

        self.label_gabarito = tk.Label(self.current_frame, text="Gabarito:", bg="#333333", fg="#FFFFFF")
        self.label_gabarito.grid(row=2, column=0, padx=5, pady=5, sticky="w")
        self.entry_gabarito = tk.Entry(self.current_frame, width=10, bg="#555555", fg="#FFFFFF", insertbackground="#FFFFFF")
        self.entry_gabarito.grid(row=2, column=1, padx=5, pady=5, sticky="ew")

        self.btn_anexar_imagem = tk.Button(self.current_frame, text="Anexar Imagem", command=self.anexar_imagem, bg="#007BFF", fg="#FFFFFF")
        self.btn_anexar_imagem.grid(row=3, column=0, padx=5, pady=5, sticky="ew")

        self.image_path = None
        self.images = {}  # Dicionário para armazenar imagens

        self.btn_adicionar = tk.Button(self.current_frame, text="Adicionar Questão", command=self.adicionar_questao, bg="#28a745", fg="#FFFFFF")
        self.btn_adicionar.grid(row=3, column=1, padx=5, pady=10, sticky="e")

        self.btn_recomecar = tk.Button(self.current_frame, text="Recomeçar", command=self.recomecar, bg="#FF5733", fg="#FFFFFF")
        self.btn_recomecar.grid(row=4, column=0, padx=5, pady=10, sticky="w")

        self.btn_gerar = tk.Button(self.current_frame, text="Gerar Simulados", command=self.gerar_simulados, bg="#17a2b8", fg="#FFFFFF")
        self.btn_gerar.grid(row=4, column=1, padx=5, pady=10, sticky="e")

        self.btn_view_questions = tk.Button(self.current_frame, text="Ver Questões", command=self.show_questions_frame, bg="#17a2b8", fg="#FFFFFF")
        self.btn_view_questions.grid(row=5, column=0, columnspan=2, pady=10)

        # Botão para abrir o Leitor de Cartão
        self.btn_abrir_leitor = tk.Button(self.current_frame, text="Ler Cartão Resposta",
                                         command=self.abrir_leitor_cartao, bg="#17a2b8", fg="#FFFFFF")
        self.btn_abrir_leitor.grid(row=6, column=0, columnspan=2, pady=10)

    def abrir_leitor_cartao(self):
        """Abre a interface do leitor de cartão."""
        self.root.withdraw()
        leitor_cartao.iniciar_leitor(self.root)

    def show_questions_frame(self):
        """Exibe o quadro com a lista de questões adicionadas."""
        if self.current_frame:
            self.current_frame.destroy()

        self.current_frame = tk.Frame(self.root, bg="#333333")
        self.current_frame.pack(pady=20, expand=True, fill=tk.BOTH)

        self.listbox = tk.Listbox(self.current_frame, width=80, height=20, bg="#555555", fg="#FFFFFF")
        self.listbox.pack(padx=10, pady=10)

        self.btn_edit = tk.Button(self.current_frame, text="Editar Questão", command=self.edit_question, bg="#28a745", fg="#FFFFFF")
        self.btn_edit.pack(pady=5)

        self.btn_delete = tk.Button(self.current_frame, text="Excluir Questão", command=self.delete_question, bg="#FF5733", fg="#FFFFFF")
        self.btn_delete.pack(pady=5)

        self.btn_back = tk.Button(self.current_frame, text="Voltar", command=self.show_entry_frame, bg="#007BFF", fg="#FFFFFF")
        self.btn_back.pack(pady=5)

        self.update_question_list()

    def update_question_list(self):
        """Atualiza a lista de questões exibidas no quadro de questões."""
        self.listbox.delete(0, tk.END)
        for questao in self.questoes:
            self.listbox.insert(tk.END, questao['Enunciado'])  # Exibe a descrição da questão

    def edit_question(self):
        """Carrega a questão selecionada nos campos de entrada para edição."""
        selected_index = self.listbox.curselection()
        if not selected_index:
            messagebox.showwarning("Aviso", "Selecione uma questão para editar.")
            return

        index = selected_index[0]
        questao = self.questoes[index]

        # Volta para o quadro de entrada
        self.show_entry_frame()

        # Preenche os campos de entrada com os detalhes da questão selecionada
        self.text_enunciado.delete("1.0", tk.END)
        self.text_enunciado.insert(tk.END, questao['Enunciado'])
        self.entry_disciplina.delete(0, tk.END)
        self.entry_disciplina.insert(0, questao['Disciplina'])
        self.entry_gabarito.delete(0, tk.END)
        self.entry_gabarito.insert(0, questao['Gabarito'])

        # Limpa as imagens atuais e reinsere as imagens da questão selecionada
        self.images = questao['Imagens'].copy()  # Carrega imagens para edição
        for img_id, img_data in self.images.items():
            img_bytes = base64.b64decode(img_data)
            img_io = BytesIO(img_bytes)
            photo = ImageTk.PhotoImage(Image.open(img_io))
            self.text_enunciado.image_create(tk.END, image=photo)

            # Armazena a referência para evitar a coleta de lixo
            if not hasattr(self, 'image_refs'):
                self.image_refs = []
            self.image_refs.append(photo)

        # Define o índice de edição
        self.editing_index = index

    def delete_question(self):
        """Exclui a questão selecionada da lista."""
        selected_index = self.listbox.curselection()
        if not selected_index:
            messagebox.showwarning("Aviso", "Selecione uma questão para excluir.")
            return

        index = selected_index[0]
        del self.questoes[index]
        self.update_question_list()
        messagebox.showinfo("Sucesso", "Questão excluída com sucesso!")

    def anexar_imagem(self):
        """Permite ao usuário anexar uma imagem à questão."""
        file_path = filedialog.askopenfilename(title="Selecione uma Imagem", filetypes=[("Imagem", "*.png *.jpg *.jpeg *.gif")])
        if file_path:
            try:
                img = Image.open(file_path)
                img.thumbnail((300, 300))  # Ajuste o tamanho da miniatura conforme necessário
                photo = ImageTk.PhotoImage(img)
                self.text_enunciado.image_create(tk.INSERT, image=photo)

                # Guardar referência à imagem para evitar que seja removida pelo garbage collector
                if not hasattr(self, 'image_refs'):
                    self.image_refs = []
                self.image_refs.append(photo)

                # Salvar a imagem como base64 para posterior inserção no documento Word
                buffered = BytesIO()
                img.save(buffered, format="PNG")
                img_str = base64.b64encode(buffered.getvalue()).decode()

                img_id = str(uuid.uuid4())
                self.images[img_id] = img_str
                self.text_enunciado.insert(tk.INSERT, f"<img:{img_id}>")

                messagebox.showinfo("Sucesso", "Imagem anexada com sucesso!")
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao anexar a imagem: {e}")

    def paste_text_with_image(self, event=None):
        """Permite colar texto com imagens no widget de texto."""
        try:
            clipboard_data = self.root.clipboard_get()
            self.text_enunciado.insert(tk.INSERT, clipboard_data)
        except tk.TclError:
            try:
                image = ImageGrab.grabclipboard()
                if image:
                    image.thumbnail((300, 300))  # Ajuste o tamanho da miniatura conforme necessário
                    photo = ImageTk.PhotoImage(image)
                    self.text_enunciado.image_create(tk.INSERT, image=photo)

                    # Guardar referência à imagem para evitar que seja removida pelo garbage collector
                    if not hasattr(self, 'image_refs'):
                        self.image_refs = []
                    self.image_refs.append(photo)

                    # Salvar a imagem como base64 para posterior inserção no documento Word
                    buffered = BytesIO()
                    image.save(buffered, format="PNG")
                    img_str = base64.b64encode(buffered.getvalue()).decode()

                    img_id = str(uuid.uuid4())
                    self.images[img_id] = img_str
                    self.text_enunciado.insert(tk.INSERT, f"<img:{img_id}>")

                    messagebox.showinfo("Sucesso", "Imagem colada com sucesso!")
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao colar a imagem: {e}")
        return "break"

    def adicionar_questao(self):
        """Adiciona uma nova questão à lista de questões."""
        enunciado = self.text_enunciado.get("1.0", tk.END).strip()
        disciplina = self.entry_disciplina.get().strip()
        gabarito = self.entry_gabarito.get().strip()

        if not enunciado or not disciplina or not gabarito:
            messagebox.showwarning("Aviso", "Todos os campos devem ser preenchidos.")
            return

        # Verifica se estamos editando uma questão existente
        if self.editing_index is not None:
            self.questoes[self.editing_index] = {'Enunciado': enunciado, 'Disciplina': disciplina, 'Gabarito': gabarito, 'Imagens': self.images.copy()}
            self.editing_index = None  # Limpa o índice de edição
            messagebox.showinfo("Sucesso", "Questão editada com sucesso!")
        else:
            self.questoes.append({'Enunciado': enunciado, 'Disciplina': disciplina, 'Gabarito': gabarito, 'Imagens': self.images.copy()})
            messagebox.showinfo("Sucesso", "Questão adicionada com sucesso!")

        self.text_enunciado.delete("1.0", tk.END)
        self.entry_disciplina.delete(0, tk.END)
        self.entry_gabarito.delete(0, tk.END)
        self.images.clear()
        self.update_question_list()

    def recomecar(self):
        """Limpa todas as questões adicionadas."""
        self.questoes.clear()
        self.text_enunciado.delete("1.0", tk.END)
        self.entry_disciplina.delete(0, tk.END)
        self.entry_gabarito.delete(0, tk.END)
        self.images.clear()
        messagebox.showinfo("Sucesso", "Todas as questões foram removidas. Você pode recomeçar.")

    def gerar_simulados(self):
        """Gera os simulados em arquivos Word e Excel."""
        if not self.questoes:
            messagebox.showwarning("Aviso", "Nenhuma questão adicionada.")
            return

        dir_path = filedialog.askdirectory(title="Selecione o Diretório para Salvar os Arquivos")
        if not dir_path:
            return

        df = pd.DataFrame(self.questoes)
        df['Indice'] = df.index.map(lambda x: f"QUESTÃO {x + 1:02d}")  # Formata como QUESTÃO 01, QUESTÃO 02, etc.

        cadernos = ['Padrão', 'Azul', 'Amarelo', 'Cinza', 'Rosa']
        embaralhados = {'Padrão': df}

        for caderno in cadernos[1:]:
            embaralhados[caderno] = df.sample(frac=1).reset_index(drop=True)

        try:
            self.gerar_word(embaralhados, dir_path)
            self.gerar_excel(df, embaralhados, dir_path)
            messagebox.showinfo("Sucesso", "Documentos Word e planilha Excel gerados com sucesso!")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar documentos: {e}")

    def gerar_word(self, embaralhados, dir_path):
        """Gera os arquivos Word dos simulados."""
        pattern = re.compile(r'([A-Ea-e]\))')

        for caderno, df in embaralhados.items():
            doc = Document()
            doc.add_heading(f"Simulado {caderno}", 0)

            section = doc.sections[0]

            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

            sectPr = section._sectPr
            cols = OxmlElement('w:cols')
            cols.set(qn('w:num'), '2')
            cols.set(qn('w:sep'), '1')
            sectPr.append(cols)

            for i, (index, row) in enumerate(df.iterrows()):
                p = doc.add_heading(f"QUESTÃO {index + 1:02d} ", level=1)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.bold = True
                    run.underline = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)

                # Cria a imagem do retângulo
                largura_retangulo = int(1.80* 96)  # Convertendo de polegadas para pixels (5.5 polegadas * 96 pixels por polegada)
                altura_retangulo = int(0.08 * 96)   # Convertendo de polegadas para pixels (0.1 polegadas * 96 pixels por polegada)
                imagem_retangulo = Image.new('RGB', (largura_retangulo, altura_retangulo), color='black')

                # Salva a imagem temporariamente em um buffer
                buffered = BytesIO()
                imagem_retangulo.save(buffered, format="PNG")

                # Adiciona a imagem ao documento diretamente do buffer
                run.add_picture(buffered)

                # Limpa o buffer após o uso
                buffered.close()

                # Adiciona o conteúdo da questão
                enunciado_parts = row['Enunciado'].split("<img:")
                for part in enunciado_parts:
                    if ">" in part:
                        img_id, text = part.split(">", 1)
                        if img_id in row['Imagens']:
                            img_data = base64.b64decode(row['Imagens'][img_id])
                            img_io = BytesIO(img_data)

                            # Redimensiona a imagem se ela exceder a largura máxima
                            max_image_width = Inches(2.5)

                            img = Image.open(img_io)
                            if img.width > max_image_width * 96:  # 96 pixels por polegada
                                img.thumbnail((max_image_width * 96, img.height * (max_image_width * 96 / img.width)))

                            img_io_resized = BytesIO()
                            img.save(img_io_resized, format="PNG")
                            img_io_resized.seek(0)  # Reinicia a posição do buffer

                            doc.add_picture(img_io, width=max_image_width)

                        p = doc.add_paragraph()
                        self.adicionar_texto_com_formatacao(p, text, pattern)
                    else:
                        p = doc.add_paragraph()
                        self.adicionar_texto_com_formatacao(p, part, pattern)

                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(14)
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.right_indent = Inches(0)

            doc.save(f"{dir_path}/Simulado_{caderno}.docx")

    def adicionar_texto_com_formatacao(self, paragraph, text, pattern):
        """Aplica formatação específica a partes do texto."""
        parts = pattern.split(text)
        for part in parts:
            if pattern.match(part):
                run = paragraph.add_run(part[0].upper())
                run.font.name = '123Testing2'
                run.font.size = Pt(12)
            else:
                run = paragraph.add_run(part)
                run.font.name = 'Arial'
                run.font.size = Pt(11)

    def gerar_excel(self, original_df, embaralhados, dir_path):
        """Gera a planilha Excel com a correlação das questões."""
        correlacao = pd.DataFrame(columns=['Disciplina', 'Gabarito', 'Padrão', 'Azul', 'Amarelo', 'Cinza', 'Rosa'])

        for i, questao in original_df.iterrows():
            correlacao.loc[i] = [
                questao['Disciplina'],
                questao['Gabarito'],
                embaralhados['Padrão'].loc[embaralhados['Padrão']['Indice'] == questao['Indice']].index[0] + 1,
                embaralhados['Azul'].loc[embaralhados['Azul']['Indice'] == questao['Indice']].index[0] + 1,
                embaralhados['Amarelo'].loc[embaralhados['Amarelo']['Indice'] == questao['Indice']].index[0] + 1,
                embaralhados['Cinza'].loc[embaralhados['Cinza']['Indice'] == questao['Indice']].index[0] + 1,
                embaralhados['Rosa'].loc[embaralhados['Rosa']['Indice'] == questao['Indice']].index[0] + 1
            ]

        try:
            correlacao.to_excel(f"{dir_path}/correlacao_questoes.xlsx", index=False)
            print("Arquivo Excel salvo com sucesso.")
        except Exception as e:
            print(f"Erro ao salvar arquivo Excel: {e}")

def main():
    root = tk.Tk()
    app = SimuladoApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
